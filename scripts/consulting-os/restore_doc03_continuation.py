#!/usr/bin/env python3
"""Replace the stale post-Section-11 tail of Canonical Document 03.

The continuation is an explicit human-provided canonical amendment. This tool
does not infer content: it preserves the existing Sections 1-11 and replaces
the old Section 26-to-end tail with the supplied Section 12-to-end Markdown.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SECTION_TAIL_MARKER = "26. Minimum Domain Acceptance Tests"
MARKDOWN_TAIL_MARKER = "# 26. Minimum Domain Acceptance Tests"


def body_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t"))).strip()


def remove_existing_tail(document: Document) -> None:
    body = document.element.body
    children = list(body.iterchildren())
    start_index = None
    for index, child in enumerate(children):
        if child.tag == qn("w:p") and body_text(child) == SECTION_TAIL_MARKER:
            start_index = index
            break
    if start_index is None:
        raise ValueError(f"Could not find DOCX tail marker: {SECTION_TAIL_MARKER}")

    for child in children[start_index:]:
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_inline_text(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table_widths(headers: list[str], total: int) -> list[int]:
    count = len(headers)
    first = headers[0].strip().lower() if headers else ""
    if count == 2:
        ratios = [0.30, 0.70]
    elif count == 3 and first == "relationship":
        ratios = [0.20, 0.34, 0.46]
    elif count == 3 and first == "step":
        ratios = [0.08, 0.23, 0.69]
    elif count == 3:
        ratios = [0.22, 0.39, 0.39]
    elif count == 4:
        ratios = [0.14, 0.24, 0.39, 0.23]
    else:
        ratios = [1 / count] * count
    widths = [int(total * ratio) for ratio in ratios]
    widths[-1] += total - sum(widths)
    return widths


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    normalized = [row + [""] * (columns - len(row)) for row in rows]
    table = document.add_table(rows=len(normalized), cols=columns)
    table.style = "Table Grid"
    table.autofit = False

    section = document.sections[-1]
    usable_twips = int((section.page_width - section.left_margin - section.right_margin) / 635)
    widths = table_widths(normalized[0], usable_twips)

    table_pr = table._tbl.tblPr
    table_width = table_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(usable_twips))
    table_width.set(qn("w:type"), "dxa")

    table_indent = table_pr.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_pr.append(table_indent)
    # Align the visible table border with body text after Word's default
    # 120-DXA start-cell margin.
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")

    table_grid = table._tbl.tblGrid
    for child in list(table_grid):
        table_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        table_grid.append(grid_col)

    for row_index, row_values in enumerate(normalized):
        for column_index, value in enumerate(row_values):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            add_inline_text(paragraph, value)
            set_cell_width(cell, widths[column_index])
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
        if row_index == 0:
            tr_pr = table.rows[row_index]._tr.get_or_add_trPr()
            table_header = OxmlElement("w:tblHeader")
            table_header.set(qn("w:val"), "true")
            tr_pr.append(table_header)


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def append_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line.strip():
            index += 1
            continue

        if line.startswith("|") and index + 1 < len(lines) and re.match(
            r"^\|(?:\s*:?-+:?\s*\|)+$", lines[index + 1].strip()
        ):
            rows = [parse_table_row(line)]
            index += 2
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                rows.append(parse_table_row(lines[index]))
                index += 1
            add_table(document, rows)
            continue

        heading = re.match(r"^(#{1,2})\s+(.+)$", line)
        if heading:
            paragraph = document.add_paragraph(style=f"Heading {len(heading.group(1))}")
            paragraph.paragraph_format.keep_with_next = True
            add_inline_text(paragraph, heading.group(2))
            index += 1
            continue

        nested_bullet = re.match(r"^\s{2,}-\s+(.+)$", line)
        if nested_bullet:
            style = "List Bullet 2" if "List Bullet 2" in document.styles else "List Bullet"
            paragraph = document.add_paragraph(style=style)
            if style == "List Bullet":
                paragraph.paragraph_format.left_indent = Inches(0.5)
            add_inline_text(paragraph, nested_bullet.group(1))
            index += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", line)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_text(paragraph, bullet.group(1))
            index += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            add_inline_text(paragraph, numbered.group(1))
            index += 1
            continue

        paragraph = document.add_paragraph()
        if line.startswith("   "):
            paragraph.paragraph_format.left_indent = Inches(0.5)
        add_inline_text(paragraph, line.strip())
        index += 1


def build_markdown(source_markdown: str, continuation: str) -> str:
    if MARKDOWN_TAIL_MARKER not in source_markdown:
        raise ValueError(f"Could not find Markdown tail marker: {MARKDOWN_TAIL_MARKER}")
    prefix = source_markdown.split(MARKDOWN_TAIL_MARKER, 1)[0].rstrip()
    merged = prefix + "\n\n" + continuation.strip() + "\n"
    for number in range(12, 31):
        marker = rf"^# {number}\."
        count = len(re.findall(marker, merged, flags=re.MULTILINE))
        if count != 1:
            raise ValueError(f"Expected exactly one level-1 Section {number} heading; found {count}")
    if merged.count("# Appendix A — Canonical Entity Index") != 1:
        raise ValueError("Expected exactly one Appendix A heading.")
    if not merged.rstrip().endswith("# End of Canonical Document 03"):
        raise ValueError("Continuation does not end with the canonical end marker.")
    return merged


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-docx", type=Path, required=True)
    parser.add_argument("--source-markdown", type=Path, required=True)
    parser.add_argument("--continuation", type=Path, required=True)
    parser.add_argument("--output-docx", type=Path, required=True)
    parser.add_argument("--output-markdown", type=Path, required=True)
    args = parser.parse_args()

    continuation = args.continuation.read_text(encoding="utf-8")
    source_markdown = args.source_markdown.read_text(encoding="utf-8")
    merged_markdown = build_markdown(source_markdown, continuation)

    document = Document(args.source_docx)
    remove_existing_tail(document)
    append_markdown(document, continuation)

    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    args.output_markdown.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output_docx)
    args.output_markdown.write_text(merged_markdown, encoding="utf-8")


if __name__ == "__main__":
    main()
